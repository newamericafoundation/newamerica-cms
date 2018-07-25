from docx import Document

from .section import Section
from .parsers import parse_endnotes, parse_links

class DocxParse():
    def __init__(self, ref):
        self._ref = ref
        self._doc = Document(self._ref)
        self._links = parse_links(self._doc.part.rels)
        self.endnotes = parse_endnotes(self._doc.part.package.parts)
        self.sections = self.__sections__()

    def __sections__(self):
        # group all runs for a given section
        # then compile blocks for each section
        sections = []
        section = None

        for p in self._doc.paragraphs:
            if p.style.name == 'Heading 1':
                section = Section(p, self._doc, self._links, self.endnotes)
                sections.append(section)
                continue
            # all blocks must have a section
            if section:
                section.add_paragraph(p)

        for s in sections:
            s.compile_blocks()

        return sections
